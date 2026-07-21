"""
Document loaders.

Supports PDF, TXT, Markdown, and DOCX transparently: callers just point
`load_documents` at a directory and get back a uniform list of
`RawDocument` objects, regardless of source file type.
"""

from __future__ import annotations

from pathlib import Path

from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from ai_core.rag.logging_utils import get_logger
from ai_core.rag.models import DocumentType, RawDocument

logger = get_logger(__name__)

_EXTENSION_TO_TYPE = {
    ".pdf": DocumentType.PDF,
    ".txt": DocumentType.TXT,
    ".md": DocumentType.MARKDOWN,
    ".docx": DocumentType.DOCX,
}


class UnsupportedDocumentTypeError(Exception):
    """Raised when a document with an unsupported extension is encountered."""


def _load_pdf(path: Path) -> list[RawDocument]:
    """Load a PDF, producing one RawDocument per page (page numbers preserved)."""
    documents: list[RawDocument] = []
    reader = PdfReader(str(path))
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        documents.append(
            RawDocument(
                filename=path.name,
                source=str(path),
                document_type=DocumentType.PDF,
                text=text,
                page=page_number,
            )
        )
    return documents


def _load_txt(path: Path) -> list[RawDocument]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return [
        RawDocument(
            filename=path.name,
            source=str(path),
            document_type=DocumentType.TXT,
            text=text,
            page=None,
        )
    ]


def _load_markdown(path: Path) -> list[RawDocument]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return [
        RawDocument(
            filename=path.name,
            source=str(path),
            document_type=DocumentType.MARKDOWN,
            text=text,
            page=None,
        )
    ]


def _load_docx(path: Path) -> list[RawDocument]:
    doc = DocxDocument(str(path))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    return [
        RawDocument(
            filename=path.name,
            source=str(path),
            document_type=DocumentType.DOCX,
            text=text,
            page=None,
        )
    ]


_LOADERS = {
    ".pdf": _load_pdf,
    ".txt": _load_txt,
    ".md": _load_markdown,
    ".docx": _load_docx,
}


def load_document(path: Path) -> list[RawDocument]:
    """Load a single document file into one or more RawDocument objects."""
    suffix = path.suffix.lower()
    loader = _LOADERS.get(suffix)
    if loader is None:
        raise UnsupportedDocumentTypeError(f"Unsupported document type: {suffix}")

    logger.info("Loading document: %s", path)
    try:
        return loader(path)
    except Exception:
        logger.exception("Failed to load document: %s", path)
        raise


def load_documents(directory: Path) -> list[RawDocument]:
    """Load every supported document inside a directory (non-recursive)."""
    if not directory.exists():
        logger.warning("Documents directory does not exist: %s", directory)
        return []

    all_documents: list[RawDocument] = []
    for path in sorted(directory.iterdir()):
        if not path.is_file():
            continue
        if path.suffix.lower() not in _EXTENSION_TO_TYPE:
            logger.debug("Skipping unsupported file: %s", path)
            continue
        all_documents.extend(load_document(path))

    logger.info("Loaded %d raw document segment(s) from %s", len(all_documents), directory)
    return all_documents
